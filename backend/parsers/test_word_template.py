from docx import Document

FILE = r"C:\Users\Eunece\Documents\MileageTracker\backend\templates\SARS Logbook 2026.docx"

doc = Document(FILE)

print("Document opened successfully!")
print("Tables found:", len(doc.tables))

# Show all table contents
for table_index, table in enumerate(doc.tables):
    print(f"\n--- TABLE {table_index} ---")

    for row in table.rows:
        row_data = []

        for cell in row.cells:
            text = cell.text.strip()
            row_data.append(text)

        print(row_data)