from docx import Document
import pandas as pd


EXCEL_FILE = "backend/exports/mileage_logbook.xlsx"

TEMPLATE_FILE = "backend/templates/SARS Logbook 2026.docx"

OUTPUT_FILE = "backend/exports/Completed_SARS_Logbook_2026.docx"


def export_word():

    print("Loading Excel trips...")

    df = pd.read_excel(EXCEL_FILE, sheet_name="Trips")

    print(f"Trips loaded: {len(df)}")

    # SORT
    df = df.sort_values(by="Date").reset_index(drop=True)

    print("Loading Word template...")

    doc = Document(TEMPLATE_FILE)

    print(f"Tables found: {len(doc.tables)}")

    trip_index = 0

    # ============================================
    # SARS MONTH TABLES
    # ============================================

    for table_index in range(2, len(doc.tables)):

        table = doc.tables[table_index]

        print(f"Using table {table_index}")

        # START AFTER HEADER ROWS
        for row in table.rows[2:]:

            if trip_index >= len(df):
                break

            cells = row.cells

            if len(cells) < 7:
                continue

            trip = df.iloc[trip_index]

            try:

                km = float(trip["KM"])

                # ============================================
                # SARS COLUMN MAP
                # ============================================

                # 0 = DATE
                cells[0].text = str(trip["Date"])

                # 1 = OPENING KM
                cells[1].text = ""

                # 2 = CLOSING KM
                cells[2].text = ""

                # 3 = TOTAL BUSINESS KM
                cells[3].text = str(round(km, 2))

                # 4 = FROM
                cells[4].text = str(trip["From"])

                # 5 = TO
                cells[5].text = str(trip["To"])

                # 6 = PURPOSE
                cells[6].text = "Business Travel"

                trip_index += 1

            except Exception as e:
                print("Row error:", e)

    print("\nSaving completed document...")

    doc.save(OUTPUT_FILE)

    print("\n===================================")
    print("WORD EXPORT COMPLETE")
    print("===================================")

    print(f"Trips inserted: {trip_index}")

    print(f"Saved to:\n{OUTPUT_FILE}")


def main():
    export_word()


if __name__ == "__main__":
    main()